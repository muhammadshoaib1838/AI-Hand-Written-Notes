import streamlit as st
import easyocr
import cv2
import numpy as np
from groq import Groq
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import os

# --- Page Config ---
st.set_page_config(page_title="AI Note Pro", page_icon="📝", layout="wide")

# --- Professional Styling ---
st.markdown("""
    <style>
    .stApp { background-color: #0e1117; color: #f1f5f9; }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; }
    .stTabs [data-baseweb="tab"] { 
        height: 50px; 
        background-color: #1e293b; 
        border-radius: 10px 10px 0 0; 
        color: white; 
    }
    .stTabs [aria-selected="true"] { background-color: #6366f1 !important; }
    
    /* Button Styling */
    div.stButton > button {
        background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%) !important;
        color: white !important;
        font-weight: bold !important;
        border: none !important;
        border-radius: 8px;
        width: 100%;
    }
    
    .signature { color: #818cf8; text-align: center; margin-top: 50px; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

# --- Init OCR & API ---
@st.cache_resource
def load_ocr():
    # gpu=False is required for Streamlit Cloud Free Tier
    return easyocr.Reader(['en'], gpu=False)

reader = load_ocr()
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# --- Logic Functions ---
def process_pipeline(uploaded_file):
    try:
        # 1. OCR Stage
        uploaded_file.seek(0)
        file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
        img = cv2.imdecode(file_bytes, 1)
        
        if img is None:
            return "Failed to decode image", None, None

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        processed_img = cv2.adaptiveThreshold(
            cv2.GaussianBlur(gray, (5, 5), 0),
            255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        
        raw_results = reader.readtext(processed_img, detail=0)
        raw_text = " ".join(raw_results)
        
        if not raw_text.strip():
            return "No text detected in image", None, None

        # 2. AI Structuring Stage
        if not client:
            return raw_text, "Error: Groq API Key is missing or invalid.", "N/A"

        # Explicitly joining strings to avoid unterminated literal errors
        prompt_parts = [
            "Act as a Professional Document Architect.",
            "Convert this OCR text into a beautiful digital document:",
            "- Use # for the Main Title",
            "- Use ## for Section Headings",
            "- Use bolding (**) for important technical terms",
            "- End with a section titled '--- FINAL SUMMARY ---'",
            f"TEXT: {raw_text}"
        ]
        prompt = " ".join(prompt_parts)

        completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile"
        )
        
        full_output = completion.choices[0].message.content
        
        if "--- FINAL SUMMARY ---" in full_output:
            blueprint, summary = full_output.split("--- FINAL SUMMARY ---")
        else:
            blueprint, summary = full_output, "Summary not explicitly separated by AI."
            
        return raw_text, blueprint.strip(), summary.strip()
    except Exception as e:
        return f"Pipeline Error: {str(e)}", None, None

# --- Main UI ---
st.title("🌌 Document Architect Pro")
st.write("Upload an image of notes to transform them into structured documents.")

if not GROQ_API_KEY:
    st.error("🔑 Groq API Key not found. Please add it to your Streamlit Secrets.")

uploaded_file = st.file_uploader("📸 Choose an image...", type=["jpg", "jpeg", "png"])

if uploaded_file is not None:
    if st.button("ARCHITECT DOCUMENT ✨"):
        with st.spinner("Processing document architecture..."):
            raw, notes, summary = process_pipeline(uploaded_file)
            
            if notes:
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    t1, t2, t3 = st.tabs(["✨ Digital Blueprint", "💡 Summary", "📄 Raw OCR"])
                    
                    with t1:
                        st.markdown(notes)
                        # Word Export
                        doc = Document()
                        doc.add_heading('Architected Document', 0)
                        doc.add_paragraph(notes)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        st.download_button(
                            label="Download DOCX",
                            data=doc_io.getvalue(),
                            file_name="Architect_Export.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                    with t2:
                        st.info(summary)

                    with t3:
                        st.text_area("OCR Buffer", raw, height=300)
                
                with col2:
                    st.image(uploaded_file, caption="Uploaded Source", use_container_width=True)
            else:
                st.error(f"Failed to process: {raw}")

st.markdown('<div class="signature">Designed & Engineered by Muhammad Shoaib Nazz</div>', unsafe_allow_html=True)
